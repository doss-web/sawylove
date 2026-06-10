from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# ---- Title ----
title = doc.add_heading('', level=0)
run = title.add_run('纸片人男友 — 数据库配置大白话笔记')
run.font.size = Pt(18)
run.font.bold = True

doc.add_paragraph('写给新手看的，没有术语，全是大白话。', style='Normal')
doc.add_paragraph('')

# ---- Section 1 ----
doc.add_heading('一、数据库是什么？', level=1)
doc.add_paragraph(
    '数据库就是一个超大的 Excel 表格，只不过它不是存在你电脑上，而是存在网上的服务器里。\n\n'
    '你聊天时填的表单、发的消息、AI 的记忆，全存在这张"网上表格"里。'
)

# ---- Section 2 ----
doc.add_heading('二、为什么不用自己电脑存？', level=1)
doc.add_paragraph(
    '因为你的聊天应用最终要开放给全世界用。如果数据存你电脑上，别人打开这个网站就看不到自己的聊天记录。\n'
    '必须存在"公共储藏室"——就是云数据库。'
)

# ---- Section 3 ----
doc.add_heading('三、为什么选 Supabase？', level=1)
doc.add_paragraph(
    '市面上有很多云数据库服务商，Supabase 就像一家信誉好的仓库出租公司，重点是：'
)
bullet = doc.add_paragraph('', style='List Bullet')
bullet.add_run('有免费仓库').bold = True
bullet.add_run(' — 500MB 够我们用了')

bullet2 = doc.add_paragraph('', style='List Bullet')
bullet2.add_run('仓库在东京').bold = True
bullet2.add_run(' — 离国内近，网速快')

bullet3 = doc.add_paragraph('', style='List Bullet')
bullet3.add_run('配了管理员（Pooler）').bold = True
bullet3.add_run(' — 这个最关键，下面细说')

# ---- Section 4 ----
doc.add_heading('四、Pooler 是什么？用快递站来类比', level=1)
doc.add_paragraph(
    '普通的数据库连接就像打电话：\n\n'
    '    你的网站 → 拨号 → 数据库接听 → 挂断\n'
    '    每次有人来聊天就拨一次\n\n'
    '人多的时候，数据库要接无数个电话，直接崩了。\n\n'
    'Pooler 就像快递站的收发员：\n\n'
    '    你的网站 → 把"查数据"的需求丢给收发员 →\n'
    '    收发员统一帮你去仓库取 → 返回结果\n\n'
    '你不用自己跑一趟仓库，收发员帮你管理"哪些查完了、哪些还在排队"，效率高很多。\n\n'
    'Supabase 免费送了收发员服务，所以我们选 Session Pooler（5432 端口），意思就是"走收发员通道"。'
)

# ---- Section 5 ----
doc.add_heading('五、连接地址逐段翻译', level=1)
doc.add_paragraph('这是我们的连接地址长什么样：', style='Normal')

# Add as code block
code = doc.add_paragraph('')
code.style = doc.styles['Normal']
run = code.add_run(
    'postgresql://postgres.qzxrpwowgrchfghyzebn:密码@aws-1-ap-northeast-1.pooler.supabase.com:5432/postgres'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')

# Create table
table = doc.add_table(rows=8, cols=2, style='Table Grid')
table.autofit = True

rows_data = [
    ('这一段', '翻译成人话'),
    (': ', ''),
    ('postgresql://', '我要连 PostgreSQL 类型的仓库'),
    ('postgres.qzxrp...', '账号是 postgres，仓库编号是 qzxrpwowgrchfghyzebn'),
    (':密码', '你创建仓库时自己设的密码'),
    ('@aws-1-ap-northeast...', '@东京一号仓库的地址'),
    (':5432', '从收发员通道进去，门牌号 5432'),
    ('/postgres', '进仓库后去 "postgres" 这个区'),
]

for i, (a, b) in enumerate(rows_data):
    if i == 1:  # skip separator row
        continue
    cells = table.rows[i if i < 2 else i-1].cells
    cells[0].text = a
    cells[1].text = b
    # Make first column bold and monospace
    cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
    cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)

# Remove the separator row
# Actually let me rebuild this more simply
doc.add_paragraph('')

# Simple bullet list for each part
parts = [
    ('postgresql://', '表示"我要连的是 PostgreSQL 类型仓库"'),
    ('postgres.qzxr...', '账号 postgres，仓库编号就是那串 qzxrp...'),
    (':密码', '你创建仓库时自己设的，URL 里 / 要写 %2F'),
    ('@aws-1-ap-northeast...', '仓库地址在东京一号'),
    (':5432', '走收发员（Pooler）通道，门牌号 5432'),
    ('/postgres', '进去后去 postgres 区（默认名字）'),
]

for label, meaning in parts:
    p = doc.add_paragraph('')
    run = p.add_run(label)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.bold = True
    run = p.add_run(' → ' + meaning)
    run.font.size = Pt(11)

# ---- Section 6 ----
doc.add_heading('六、Supabase "Connect" 按钮是来干嘛的？', level=1)
doc.add_paragraph(
    '它就是仓库管理员直接把钥匙交给你。不需要你从几十个菜单里翻找，'
    '一键弹出来、直接复制整条地址，贴到代码里就完事了。\n\n'
    '以后每次忘了怎么拼地址，就去你 Supabase 项目页面顶部点那个紫色按钮。'
)

# ---- Section 7 ----
doc.add_heading('七、我们做了什么事？（五步）', level=1)

steps = [
    ('第1步', 'Supabase 官网开了一个免费仓库', '点了几下鼠标，仓库就开好了'),
    ('第2步', '复制钥匙串', '就是那条长长的连接地址'),
    ('第3步', '贴到项目的 .env 文件', '.env 就是"环境配置"文件，存各种密码和地址'),
    ('第4步', '运行 prisma db push', '等于拿钥匙试一下门能开吗'),
    ('第5步', '门开了 ✅', '数据库连上了！'),
]

for i, (step, action, detail) in enumerate(steps):
    p = doc.add_paragraph('')
    run = p.add_run(step + '：')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(action + ' — ' + detail)

# ---- Section 8 ----
doc.add_heading('八、一句话总结', level=1)
doc.add_paragraph(
    '.env 里的 DATABASE_URL 就是一把钥匙。\n'
    'Prisma 用它去开仓库门（建表），\n'
    '你的网站也用它去仓库里存取数据（聊天记录、用户信息）。\n'
    '一把钥匙，两把用途。'
)

# Save
output = 'C:\\Users\\Admin\\Desktop\\数据库配置大白话笔记.docx'
doc.save(output)
print(f'Done: {output}')

# Now delete the first table we created (it's messy) — actually the table is fine, let me remove it.
# Re-save without the ugly table
